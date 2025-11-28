"""Simple converter: Markdown -> DOCX

Reads a Markdown file (limited formatting) and writes a Word document (.docx).
Requires: python-docx

Usage:
    python tools/generate_word_doc.py input.md output.docx

This script supports headings (#, ##, ###), paragraphs and simple lists.
"""
import sys
from pathlib import Path
from docx import Document

HEADING_MAP = {1: 'HEADING_1', 2: 'HEADING_2', 3: 'HEADING_3'}


def add_paragraph_from_markdown(doc: Document, line: str):
    stripped = line.lstrip()
    if stripped.startswith('#'):
        # Count level
        level = len(line) - len(line.lstrip('#'))
        text = line[level:].strip()
        if level == 1:
            doc.add_heading(text, level=1)
        elif level == 2:
            doc.add_heading(text, level=2)
        elif level == 3:
            doc.add_heading(text, level=3)
        else:
            doc.add_paragraph(text)
    elif stripped.startswith('- '):
        # simple list item
        doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
    elif stripped.startswith('```'):
        # code fence start/end - handled by toggling global state in caller
        return
    else:
        doc.add_paragraph(line.rstrip())


def convert_md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    in_code = False
    if not md_path.exists():
        raise FileNotFoundError(f"Input markdown not found: {md_path}")

    with md_path.open('r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            if line.strip().startswith('```'):
                in_code = not in_code
                if in_code:
                    # start code block - add a heading or label for clarity
                    doc.add_paragraph('Código:', style='Intense Quote')
                continue
            if in_code:
                # code block line
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Courier New'
            else:
                if line.strip() == '':
                    # blank line -> paragraph separator
                    continue
                add_paragraph_from_markdown(doc, line)

    doc.save(str(docx_path))


def main(argv):
    if len(argv) < 3:
        print('Usage: python tools/generate_word_doc.py input.md output.docx')
        return 1
    md = Path(argv[1])
    out = Path(argv[2])
    convert_md_to_docx(md, out)
    print(f'Wrote {out}')
    return 0


if __name__ == '__main__':
    raise SystemExit(main(sys.argv))
